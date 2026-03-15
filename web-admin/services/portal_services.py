from __future__ import annotations

import sqlite3
import hashlib
import shutil
import json
import os
import time
import tempfile
import zipfile
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any
from urllib import request as urlrequest
from urllib import error as urlerror
from urllib.parse import urlencode, quote

from openpyxl import load_workbook
from num2words import num2words


BASE = Path("/root/webadminbots/Infinity Projects")
APP_ROOT = Path(__file__).resolve().parent.parent
ORDER_DOC_TEMPLATES_DIR = APP_ROOT / "templates" / "order_docs"


def _ensure_db_file(db_path: Path) -> None:
    db_path.parent.mkdir(parents=True, exist_ok=True)
    if not db_path.exists():
        with sqlite3.connect(db_path):
            pass


def db_paths() -> dict[str, Path]:
    return {
        "meeting": BASE / "Meeting-booking-bot" / "meeting_bot.db",
        "broker": BASE / "Broker-booking-bot" / "broker_booking.db",
        "order": BASE / "order-bot" / "storage" / "orders.db",
        "contracts": BASE / "contract-register" / "contracts.db",
        "docflow": BASE / "doc-flow-bot" / "app" / "database.db",
    }


def _query(db_path: Path, sql: str, params: tuple[Any, ...] = ()) -> list[sqlite3.Row]:
    if not db_path.exists():
        return []
    with sqlite3.connect(db_path) as conn:
        conn.row_factory = sqlite3.Row
        return conn.execute(sql, params).fetchall()


def _execute(db_path: Path, sql: str, params: tuple[Any, ...] = ()) -> int:
    if not db_path.exists():
        return 0
    with sqlite3.connect(db_path) as conn:
        cur = conn.execute(sql, params)
        conn.commit()
        return cur.rowcount


def _table_columns(db_path: Path, table: str) -> set[str]:
    rows = _query(db_path, f"PRAGMA table_info({table})")
    return {str(row["name"]) for row in rows}


def _password_hash(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def _ensure_web_users_table(db_path: Path) -> None:
    _ensure_db_file(db_path)
    with sqlite3.connect(db_path) as conn:
        conn.execute(
        """
        CREATE TABLE IF NOT EXISTS web_users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            telegram_id TEXT NOT NULL UNIQUE,
            email TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """,
        )
        conn.commit()


def web_register_user(module: str, telegram_id: str, email: str, password: str) -> tuple[bool, str]:
    path = db_paths().get(module)
    if path is None:
        return False, "Неизвестный модуль"
    _ensure_web_users_table(path)
    clean_email = email.strip().lower()
    if "@" not in clean_email or len(password.strip()) < 4:
        return False, "Некорректная почта или слишком короткий пароль"
    exists = _query(path, "SELECT id FROM web_users WHERE email = ? LIMIT 1", (clean_email,))
    if exists:
        return False, "Почта уже зарегистрирована"
    _execute(
        path,
        "INSERT INTO web_users (telegram_id, email, password_hash) VALUES (?, ?, ?)",
        (str(telegram_id).strip(), clean_email, _password_hash(password)),
    )
    return True, "Веб-аккаунт создан"


def web_authenticate(module: str, email: str, password: str) -> tuple[bool, str, str | None]:
    path = db_paths().get(module)
    if path is None:
        return False, "Неизвестный модуль", None
    _ensure_web_users_table(path)
    clean_email = email.strip().lower()
    rows = _query(
        path,
        "SELECT telegram_id, password_hash FROM web_users WHERE email = ? LIMIT 1",
        (clean_email,),
    )
    if not rows:
        return False, "Пользователь с такой почтой не найден", None
    row = rows[0]
    if row["password_hash"] != _password_hash(password):
        return False, "Неверный пароль", None
    return True, "Вход выполнен", str(row["telegram_id"])


def web_reset_password(module: str, email: str, new_password: str) -> tuple[bool, str]:
    path = db_paths().get(module)
    if path is None:
        return False, "Неизвестный модуль"
    _ensure_web_users_table(path)
    clean_email = email.strip().lower()
    clean_password = new_password.strip()
    if "@" not in clean_email:
        return False, "Некорректная почта"
    if len(clean_password) < 4:
        return False, "Пароль должен быть не короче 4 символов"
    exists = _query(path, "SELECT id FROM web_users WHERE email = ? LIMIT 1", (clean_email,))
    if not exists:
        return False, "Пользователь с такой почтой не найден"
    _execute(path, "UPDATE web_users SET password_hash = ? WHERE email = ?", (_password_hash(clean_password), clean_email))
    return True, "Пароль обновлен. Войдите с новым паролем"


def web_delete_user(module: str, email: str) -> tuple[bool, str]:
    path = db_paths().get(module)
    if path is None:
        return False, "Неизвестный модуль"
    _ensure_web_users_table(path)
    clean_email = email.strip().lower()
    if "@" not in clean_email:
        return False, "Некорректная почта"
    rows = _query(path, "SELECT telegram_id FROM web_users WHERE email = ? LIMIT 1", (clean_email,))
    if not rows:
        return False, "Пользователь с такой почтой не найден"
    telegram_id = str(rows[0]["telegram_id"])
    info = _module_user_table(module)
    if info is not None:
        db, table, field = info
        cols = _table_columns(db, table)
        if field in cols:
            _execute(db, f"DELETE FROM {table} WHERE {field} = ?", (telegram_id,))
    affected = _execute(path, "DELETE FROM web_users WHERE email = ?", (clean_email,))
    return (affected > 0, "Пользователь удален" if affected > 0 else "Пользователь не найден")


def _module_user_table(module: str) -> tuple[Path, str, str] | None:
    path = db_paths().get(module)
    if path is None:
        return None
    if module in {"meeting", "broker", "order", "contracts"}:
        return path, "users", "telegram_id"
    if module == "docflow":
        return path, "users", "telegram_id"
    return None


def _telegram_exists(module: str, telegram_id: str) -> bool:
    info = _module_user_table(module)
    if info is None:
        return False
    path, table, field = info
    rows = _query(path, f"SELECT {field} FROM {table} WHERE {field} = ? LIMIT 1", (telegram_id,))
    return bool(rows)


def generate_telegram_id(module: str, email: str) -> str:
    seed = int(hashlib.sha256(f"{module}:{email.strip().lower()}".encode("utf-8")).hexdigest()[:12], 16)
    candidate = 1_000_000_000 + (seed % 8_000_000_000)
    for _ in range(1000):
        text = str(candidate)
        if not _telegram_exists(module, text):
            return text
        candidate += 97
        if candidate > 9_999_999_999:
            candidate = 1_000_000_000 + (candidate % 9_000_000_000)
    raise RuntimeError("Не удалось сгенерировать уникальный telegram_id")


# -------------------------
# Meeting booking module
# -------------------------

def meeting_register_user(telegram_id: int, full_name: str, department: str) -> tuple[bool, str]:
    role = "user"
    _execute(
        db_paths()["meeting"],
        "INSERT OR REPLACE INTO users (telegram_id, full_name, department, role) VALUES (?, ?, ?, ?)",
        (telegram_id, full_name.strip(), department.strip(), role),
    )
    return True, "Пользователь зарегистрирован"


def meeting_register_web(full_name: str, department: str, email: str, account_password: str) -> tuple[bool, str]:
    telegram_id = int(generate_telegram_id("meeting", email))
    ok, text = meeting_register_user(telegram_id, full_name, department)
    if not ok:
        return ok, text
    ok2, text2 = web_register_user("meeting", str(telegram_id), email, account_password)
    if not ok2:
        return False, text2
    return True, "Пользователь зарегистрирован"


def meeting_users() -> list[dict[str, Any]]:
    rows = _query(
        db_paths()["meeting"],
        "SELECT telegram_id, full_name, department, role FROM users ORDER BY telegram_id DESC LIMIT 300",
    )
    return [dict(r) for r in rows]


def meeting_users_with_email() -> list[dict[str, Any]]:
    rows = _query(
        db_paths()["meeting"],
        """
        SELECT u.telegram_id, COALESCE(w.email, '') AS email, u.full_name, u.department, u.role
        FROM users u
        LEFT JOIN web_users w ON w.telegram_id = u.telegram_id
        ORDER BY u.telegram_id DESC
        LIMIT 300
        """,
    )
    return [dict(r) for r in rows]


def meeting_get_user(telegram_id: int) -> dict[str, Any] | None:
    rows = _query(
        db_paths()["meeting"],
        "SELECT telegram_id, full_name, department, role FROM users WHERE telegram_id = ? LIMIT 1",
        (telegram_id,),
    )
    return dict(rows[0]) if rows else None


def meeting_rooms() -> list[dict[str, Any]]:
    db = db_paths()["meeting"]
    rows = _query(
        db,
        "SELECT id, name, capacity, is_active, room_type FROM rooms ORDER BY id ASC",
    )
    if not rows:
        defaults = [
            ("Сириус", 6, 1, "meeting"),
            ("Солнце", 6, 1, "meeting"),
            ("Арктур", 4, 1, "meeting"),
            ("Бетельгейзе", 3, 1, "meeting"),
            ("Вега", 3, 1, "meeting"),
            ("Мансарда", 100, 1, "class"),
            ("Аквариум", 20, 1, "class"),
            ("Малый класс", 20, 1, "class"),
        ]
        for name, capacity, is_active, room_type in defaults:
            _execute(
                db,
                "INSERT INTO rooms (name, capacity, is_active, room_type) VALUES (?, ?, ?, ?)",
                (name, capacity, is_active, room_type),
            )
        rows = _query(
            db,
            "SELECT id, name, capacity, is_active, room_type FROM rooms ORDER BY id ASC",
        )
    return [dict(r) for r in rows]


def meeting_create_booking(user_id: int, room_id: int, start_time: str, end_time: str, title: str) -> tuple[bool, str]:
    db = db_paths()["meeting"]
    try:
        s = datetime.fromisoformat(start_time)
        e = datetime.fromisoformat(end_time)
    except ValueError:
        return False, "Некорректный формат даты/времени"
    if e <= s:
        return False, "Время окончания должно быть позже начала"

    room = _query(db, "SELECT id, COALESCE(is_active, 1) AS is_active FROM rooms WHERE id = ? LIMIT 1", (room_id,))
    if not room:
        return False, "Переговорная не найдена"
    if int(room[0]["is_active"] or 0) != 1:
        return False, "Переговорная недоступна для бронирования"

    conflicts = _query(
        db,
        """
        SELECT id
        FROM bookings
        WHERE room_id = ?
          AND start_time < ?
          AND end_time > ?
        LIMIT 1
        """,
        (room_id, e.isoformat(sep=" "), s.isoformat(sep=" ")),
    )
    if conflicts:
        return False, "Выбранный интервал пересекается с существующим бронированием"

    cols = _table_columns(db, "bookings")
    fields = ["user_id", "room_id", "start_time", "end_time", "title"]
    values: list[Any] = [user_id, room_id, s.isoformat(sep=" "), e.isoformat(sep=" "), title.strip()]
    if "reminder_sent" in cols:
        fields.append("reminder_sent")
        values.append(0)
    sql = f"INSERT INTO bookings ({', '.join(fields)}) VALUES ({', '.join(['?'] * len(fields))})"
    _execute(db, sql, tuple(values))
    return True, "Бронирование создано"


def meeting_bookings(user_id: int | None = None, all_rows: bool = False) -> list[dict[str, Any]]:
    db = db_paths()["meeting"]
    sql = """
    SELECT b.id, b.user_id, u.full_name, r.name AS room_name, b.start_time, b.end_time, b.title
    FROM bookings b
    LEFT JOIN users u ON u.telegram_id = b.user_id
    LEFT JOIN rooms r ON r.id = b.room_id
    """
    params: tuple[Any, ...] = ()
    if not all_rows and user_id is not None:
        sql += " WHERE b.user_id = ? AND b.end_time > datetime('now')"
        params = (user_id,)
    elif not all_rows:
        sql += " WHERE b.end_time > datetime('now')"
    sql += " ORDER BY b.start_time ASC LIMIT 500"
    rows = _query(db, sql, params)
    return [dict(r) for r in rows]


def meeting_cancel_booking(booking_id: int, actor_user_id: int, actor_role: str) -> tuple[bool, str]:
    db = db_paths()["meeting"]
    row = _query(db, "SELECT user_id FROM bookings WHERE id = ? LIMIT 1", (booking_id,))
    if not row:
        return False, "Бронь не найдена"
    owner_id = int(row[0]["user_id"])
    if actor_user_id != owner_id and actor_role not in {"hr", "admin"}:
        return False, "Недостаточно прав для отмены"
    affected = _execute(db, "DELETE FROM bookings WHERE id = ?", (booking_id,))
    return (affected > 0, "Бронь отменена" if affected > 0 else "Бронь не найдена")


def meeting_update_role(telegram_id: int, role: str) -> tuple[bool, str]:
    affected = _execute(db_paths()["meeting"], "UPDATE users SET role = ? WHERE telegram_id = ?", (role, telegram_id))
    return (affected > 0, "Роль обновлена" if affected > 0 else "Пользователь не найден")


def meeting_update_role_by_email(email: str, role: str) -> tuple[bool, str]:
    clean_email = email.strip().lower()
    if clean_email == "":
        return False, "Укажите email"
    rows = _query(
        db_paths()["meeting"],
        "SELECT telegram_id FROM web_users WHERE lower(email) = lower(?) LIMIT 1",
        (clean_email,),
    )
    if not rows:
        return False, "Пользователь с таким email не найден"
    return meeting_update_role(int(rows[0]["telegram_id"]), role)


# -------------------------
# Broker booking module
# -------------------------

def _ensure_broker_schema() -> None:
    db = db_paths()["broker"]
    _ensure_db_file(db)
    with sqlite3.connect(db) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                telegram_id INTEGER PRIMARY KEY,
                full_name TEXT NOT NULL,
                department TEXT NOT NULL,
                role TEXT DEFAULT 'user'
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS rooms (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE,
                capacity INTEGER NOT NULL DEFAULT 1,
                is_active INTEGER NOT NULL DEFAULT 1,
                room_type TEXT DEFAULT 'broker'
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS bookings (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                room_id INTEGER NOT NULL,
                start_time TEXT NOT NULL,
                end_time TEXT NOT NULL,
                title TEXT NOT NULL,
                reminder_sent INTEGER DEFAULT 0,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        conn.commit()


def _ensure_docflow_schema() -> None:
    db = db_paths()["docflow"]
    _ensure_db_file(db)
    with sqlite3.connect(db) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                telegram_id TEXT UNIQUE,
                full_name TEXT NOT NULL DEFAULT '',
                department_no TEXT NOT NULL DEFAULT '',
                role TEXT NOT NULL DEFAULT 'agent',
                is_active INTEGER NOT NULL DEFAULT 0,
                is_approved INTEGER NOT NULL DEFAULT 0,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS applications (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                agent_id INTEGER,
                department_no TEXT NOT NULL DEFAULT '',
                deal_type TEXT NOT NULL DEFAULT '',
                contract_no TEXT NOT NULL DEFAULT '',
                address TEXT NOT NULL DEFAULT '',
                object_type TEXT NOT NULL DEFAULT '',
                head_name TEXT NOT NULL DEFAULT '',
                agent_name TEXT NOT NULL DEFAULT '',
                status TEXT NOT NULL DEFAULT 'CREATED',
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        try:
            conn.execute("ALTER TABLE applications ADD COLUMN department_no TEXT NOT NULL DEFAULT ''")
        except sqlite3.OperationalError:
            pass
        conn.commit()


def broker_register_user(telegram_id: int, full_name: str, department: str) -> tuple[bool, str]:
    _ensure_broker_schema()
    role = "user"
    _execute(
        db_paths()["broker"],
        "INSERT OR REPLACE INTO users (telegram_id, full_name, department, role) VALUES (?, ?, ?, ?)",
        (telegram_id, full_name.strip(), department.strip(), role),
    )
    return True, "Пользователь зарегистрирован"


def broker_register_web(full_name: str, department: str, email: str, account_password: str) -> tuple[bool, str]:
    _ensure_broker_schema()
    telegram_id = int(generate_telegram_id("broker", email))
    ok, text = broker_register_user(telegram_id, full_name, department)
    if not ok:
        return ok, text
    ok2, text2 = web_register_user("broker", str(telegram_id), email, account_password)
    if not ok2:
        return False, text2
    return True, "Пользователь зарегистрирован"


def broker_users_with_email() -> list[dict[str, Any]]:
    _ensure_broker_schema()
    rows = _query(
        db_paths()["broker"],
        """
        SELECT u.telegram_id, COALESCE(w.email, '') AS email, u.full_name, u.department, u.role
        FROM users u
        LEFT JOIN web_users w ON w.telegram_id = CAST(u.telegram_id AS TEXT)
        ORDER BY u.telegram_id DESC
        LIMIT 300
        """,
    )
    return [dict(r) for r in rows]


def broker_get_user(telegram_id: int) -> dict[str, Any] | None:
    _ensure_broker_schema()
    rows = _query(
        db_paths()["broker"],
        "SELECT telegram_id, full_name, department, role FROM users WHERE telegram_id = ? LIMIT 1",
        (telegram_id,),
    )
    return dict(rows[0]) if rows else None


def broker_rooms() -> list[dict[str, Any]]:
    _ensure_broker_schema()
    db = db_paths()["broker"]
    rows = _query(
        db,
        "SELECT id, name, capacity, is_active, room_type FROM rooms ORDER BY id ASC",
    )
    if not rows:
        defaults = [
            ("Брокер 1", 1, 1, "broker"),
            ("Брокер 2", 1, 1, "broker"),
            ("Брокер 3", 1, 1, "broker"),
            ("Брокер 4", 1, 1, "broker"),
            ("Брокер 5", 1, 1, "broker"),
            ("Брокер 6", 1, 1, "broker"),
        ]
        for name, capacity, is_active, room_type in defaults:
            _execute(
                db,
                "INSERT INTO rooms (name, capacity, is_active, room_type) VALUES (?, ?, ?, ?)",
                (name, capacity, is_active, room_type),
            )
        rows = _query(db, "SELECT id, name, capacity, is_active, room_type FROM rooms ORDER BY id ASC")
    return [dict(r) for r in rows]


def broker_create_booking(user_id: int, room_id: int, start_time: str, end_time: str, title: str) -> tuple[bool, str]:
    _ensure_broker_schema()
    db = db_paths()["broker"]
    try:
        s = datetime.fromisoformat(start_time)
        e = datetime.fromisoformat(end_time)
    except ValueError:
        return False, "Некорректный формат даты/времени"
    if e <= s:
        return False, "Время окончания должно быть позже начала"
    room = _query(db, "SELECT id, COALESCE(is_active, 1) AS is_active FROM rooms WHERE id = ? LIMIT 1", (room_id,))
    if not room:
        return False, "Ресурс не найден"
    if int(room[0]["is_active"] or 0) != 1:
        return False, "Ресурс недоступен для бронирования"
    conflicts = _query(
        db,
        """
        SELECT id
        FROM bookings
        WHERE room_id = ?
          AND start_time < ?
          AND end_time > ?
        LIMIT 1
        """,
        (room_id, e.isoformat(sep=" "), s.isoformat(sep=" ")),
    )
    if conflicts:
        return False, "Выбранный интервал пересекается с существующим бронированием"
    cols = _table_columns(db, "bookings")
    fields = ["user_id", "room_id", "start_time", "end_time", "title"]
    values: list[Any] = [user_id, room_id, s.isoformat(sep=" "), e.isoformat(sep=" "), title.strip()]
    if "reminder_sent" in cols:
        fields.append("reminder_sent")
        values.append(0)
    sql = f"INSERT INTO bookings ({', '.join(fields)}) VALUES ({', '.join(['?'] * len(fields))})"
    _execute(db, sql, tuple(values))
    return True, "Бронирование создано"


def broker_bookings(user_id: int | None = None, all_rows: bool = False) -> list[dict[str, Any]]:
    _ensure_broker_schema()
    db = db_paths()["broker"]
    sql = """
    SELECT b.id, b.user_id, u.full_name, r.name AS room_name, b.start_time, b.end_time, b.title
    FROM bookings b
    LEFT JOIN users u ON u.telegram_id = b.user_id
    LEFT JOIN rooms r ON r.id = b.room_id
    """
    params: tuple[Any, ...] = ()
    if not all_rows and user_id is not None:
        sql += " WHERE b.user_id = ? AND b.end_time > datetime('now')"
        params = (user_id,)
    elif not all_rows:
        sql += " WHERE b.end_time > datetime('now')"
    sql += " ORDER BY b.start_time ASC LIMIT 500"
    rows = _query(db, sql, params)
    return [dict(r) for r in rows]


def broker_cancel_booking(booking_id: int, actor_user_id: int, actor_role: str) -> tuple[bool, str]:
    _ensure_broker_schema()
    db = db_paths()["broker"]
    row = _query(db, "SELECT user_id FROM bookings WHERE id = ? LIMIT 1", (booking_id,))
    if not row:
        return False, "Бронь не найдена"
    owner_id = int(row[0]["user_id"])
    if actor_user_id != owner_id and actor_role not in {"hr", "admin", "head"}:
        return False, "Недостаточно прав для отмены"
    affected = _execute(db, "DELETE FROM bookings WHERE id = ?", (booking_id,))
    return (affected > 0, "Бронь отменена" if affected > 0 else "Бронь не найдена")


def broker_update_role(telegram_id: int, role: str) -> tuple[bool, str]:
    _ensure_broker_schema()
    affected = _execute(db_paths()["broker"], "UPDATE users SET role = ? WHERE telegram_id = ?", (role, telegram_id))
    return (affected > 0, "Роль обновлена" if affected > 0 else "Пользователь не найден")


def broker_update_role_by_email(email: str, role: str) -> tuple[bool, str]:
    _ensure_broker_schema()
    clean_email = email.strip().lower()
    if clean_email == "":
        return False, "Укажите email"
    rows = _query(
        db_paths()["broker"],
        "SELECT telegram_id FROM web_users WHERE lower(email) = lower(?) LIMIT 1",
        (clean_email,),
    )
    if not rows:
        return False, "Пользователь с таким email не найден"
    return broker_update_role(int(rows[0]["telegram_id"]), role)


# -------------------------
# Order module (PKO/RKO)
# -------------------------

ORDER_PASSWORD = "080323"


def order_register_user(telegram_id: int, password: str, full_name: str, department: str) -> tuple[bool, str]:
    if password != ORDER_PASSWORD:
        return False, "Неверный пароль регистрации"
    _execute(
        db_paths()["order"],
        "INSERT OR REPLACE INTO users (telegram_id, full_name, department, role) VALUES (?, ?, ?, COALESCE((SELECT role FROM users WHERE telegram_id = ?), 'user'))",
        (telegram_id, full_name.strip(), department.strip(), telegram_id),
    )
    return True, "Пользователь зарегистрирован"


def order_register_web(
    access_password: str, full_name: str, department: str, email: str, account_password: str
) -> tuple[bool, str]:
    telegram_id = int(generate_telegram_id("order", email))
    ok, text = order_register_user(telegram_id, access_password, full_name, department)
    if not ok:
        return ok, text
    ok2, text2 = web_register_user("order", str(telegram_id), email, account_password)
    if not ok2:
        return False, text2
    return True, "Пользователь зарегистрирован"


def order_get_user(telegram_id: int) -> dict[str, Any] | None:
    rows = _query(
        db_paths()["order"],
        "SELECT telegram_id, full_name, department, COALESCE(role, 'user') AS role FROM users WHERE telegram_id = ? LIMIT 1",
        (telegram_id,),
    )
    return dict(rows[0]) if rows else None


def _next_order_number(db: Path, key: str) -> int:
    current = _query(db, "SELECT value FROM counters WHERE key = ? LIMIT 1", (key,))
    if not current:
        _execute(db, "INSERT INTO counters (key, value) VALUES (?, ?)", (key, 1))
        return 1
    value = int(current[0]["value"]) + 1
    _execute(db, "UPDATE counters SET value = ?, updated_at = CURRENT_TIMESTAMP WHERE key = ?", (value, key))
    return value


def order_create_request(
    telegram_id: int,
    doc_type: str,
    order_date: str,
    full_name: str,
    basis_type: str,
    contract_number: str,
    contract_date: str,
    amount: float,
) -> tuple[bool, str]:
    db = db_paths()["order"]
    key = "pko_number" if doc_type == "ПКО" else "rko_number"
    doc_number = _next_order_number(db, key)
    _execute(
        db,
        """
        INSERT INTO orders (
            doc_type, doc_number, user_id, date, full_name, basis_type, contract_number, contract_date, amount, status
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'на рассмотрении')
        """,
        (doc_type, doc_number, telegram_id, order_date, full_name, basis_type, contract_number, contract_date, amount),
    )
    return True, f"Заявка {doc_type} создана"


def order_requests(user_id: int | None = None) -> list[dict[str, Any]]:
    db = db_paths()["order"]
    has_users = bool(_table_columns(db, "users"))
    if has_users:
        sql = """
        SELECT o.id, o.doc_type, o.doc_number, o.user_id, o.date, o.full_name, o.amount, o.status,
               COALESCE(o.comment, '') AS comment, o.created_at, COALESCE(u.department, '') AS department
        FROM orders o
        LEFT JOIN users u ON u.telegram_id = o.user_id
        """
    else:
        sql = """
        SELECT id, doc_type, doc_number, user_id, date, full_name, amount, status,
               COALESCE(comment, '') AS comment, created_at, '' AS department
        FROM orders
        """
    params: tuple[Any, ...] = ()
    if user_id is not None:
        sql += " WHERE user_id = ?"
        params = (user_id,)
    sql += " ORDER BY id DESC LIMIT 500"
    return [dict(r) for r in _query(db, sql, params)]


def order_get_request(order_id: int) -> dict[str, Any] | None:
    db = db_paths()["order"]
    rows = _query(
        db,
        """
        SELECT o.id, o.doc_type, o.doc_number, o.user_id, o.date, o.full_name, o.basis_type, o.contract_number,
               o.contract_date, o.amount, o.status, COALESCE(o.comment, '') AS comment, o.created_at,
               COALESCE(u.department, '') AS department
        FROM orders o
        LEFT JOIN users u ON u.telegram_id = o.user_id
        WHERE id = ?
        LIMIT 1
        """,
        (order_id,),
    )
    return dict(rows[0]) if rows else None


def order_document_path(order_id: int) -> Path:
    docs_dir = BASE / "order-bot" / "storage" / "generated_docs"
    docs_dir.mkdir(parents=True, exist_ok=True)
    return docs_dir / f"order_{order_id}.xlsx"


def _parse_order_date(value: Any) -> datetime:
    text = str(value or "").strip()
    if not text:
        return datetime.now()
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(text[:19], fmt)
        except ValueError:
            continue
    return datetime.now()


def _month_name_ru(dt: datetime) -> str:
    months = [
        "января",
        "февраля",
        "марта",
        "апреля",
        "мая",
        "июня",
        "июля",
        "августа",
        "сентября",
        "октября",
        "ноября",
        "декабря",
    ]
    return months[dt.month - 1]


def _amount_parts(value: Any) -> tuple[int, int]:
    amount = Decimal(str(value or 0)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    rub = int(amount)
    kop = int((amount - Decimal(rub)) * 100)
    return rub, kop


def _format_amount_ru(value: Any) -> str:
    rub, kop = _amount_parts(value)
    return f"{rub:,}".replace(",", " ") + f",{kop:02d}"


def _amount_words_ru(value: Any) -> str:
    rub, kop = _amount_parts(value)
    text = num2words(rub, lang="ru")
    return f"{text} рублей {kop:02d} копеек".capitalize()


def _order_template_path(doc_type: str) -> Path:
    if str(doc_type).strip().upper() == "ПКО":
        return ORDER_DOC_TEMPLATES_DIR / "pko_template.xlsx"
    return ORDER_DOC_TEMPLATES_DIR / "rko_template.xlsx"


def _fill_rko_template(path: Path, row: dict[str, Any]) -> None:
    wb = load_workbook(path)
    ws = wb[wb.sheetnames[0]]
    dt = _parse_order_date(row.get("date"))
    amount_text = _format_amount_ru(row.get("amount"))
    contract_number = str(row.get("contract_number") or "").strip()
    contract_date = str(row.get("contract_date") or "").strip()
    basis_type = str(row.get("basis_type") or "").strip()
    basis_line = basis_type or "Основание"
    if contract_number:
        basis_line += f" № {contract_number}"
    if contract_date:
        basis_line += f" от {contract_date}"

    ws["CC11"] = int(row.get("doc_number") or 0)
    ws["CT11"] = dt.strftime("%d.%m.%Y")
    ws["CC15"] = amount_text
    full_name = str(row.get("full_name") or "")
    ws["H17"] = full_name
    ws["A17"] = full_name
    ws["K19"] = basis_line
    ws["C32"] = dt.strftime("%d")
    ws["J32"] = _month_name_ru(dt)
    ws["AC32"] = dt.strftime("%Y")
    wb.save(path)


def _replace_sheet_values(ws: Any, replacements: dict[str, str]) -> None:
    for row_cells in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row_cells:
            value = cell.value
            if not isinstance(value, str):
                continue
            updated = value
            for old, new in replacements.items():
                if old in updated:
                    updated = updated.replace(old, new)
            if updated != value:
                cell.value = updated


def _fill_pko_template(path: Path, row: dict[str, Any]) -> None:
    wb = load_workbook(path)
    ws = wb[wb.sheetnames[0]]
    dt = _parse_order_date(row.get("date"))
    rub, kop = _amount_parts(row.get("amount"))
    amount_text = _format_amount_ru(row.get("amount"))
    amount_words = _amount_words_ru(row.get("amount"))
    contract_number = str(row.get("contract_number") or "").strip()
    contract_date = str(row.get("contract_date") or "").strip()
    basis_type = str(row.get("basis_type") or "").strip()
    basis_line_full = basis_type.strip()
    contract_line = (f"{contract_number} от {contract_date}").strip()

    ws["CX9"] = int(row.get("doc_number") or 0)
    ws["AQ13"] = int(row.get("doc_number") or 0)
    ws["BB13"] = dt.strftime("%d.%m.%Y")
    ws["CA10"] = dt.strftime("%d")
    ws["CG10"] = _month_name_ru(dt)
    ws["CW10"] = dt.strftime("%Y")
    full_name = str(row.get("full_name") or "")
    ws["G15"] = full_name
    ws["K16"] = full_name
    ws["K21"] = full_name
    ws["CF14"] = full_name
    for cell_addr, value in {
        "BW14": full_name,
        "BW15": full_name,
        "CC23": amount_text,
        "CC24": amount_words,
        "BW16": basis_line_full,
        "BW17": contract_line,
    }.items():
        try:
            ws[cell_addr] = value
        except Exception:
            pass
    ws["AO19"] = amount_text
    ws["CC19"] = rub
    ws["CY19"] = kop
    ws["K23"] = basis_line_full or basis_type
    ws["A24"] = contract_line
    ws["K28"] = amount_words
    ws["CG25"] = amount_words
    _replace_sheet_values(
        ws,
        {
            "Надеев Анатолий Алексеевич": full_name,
            "Корчагина Дарья Игоревна": full_name,
            "6061/1": contract_number or "—",
            "02.03.2026": contract_date or dt.strftime("%d.%m.%Y"),
            "Предоплата по договору № 6061/1": basis_line_full or basis_type,
            "Предоплата по договору №6061/1": basis_line_full or basis_type,
            "Двести семьдесят семь тысяч пятьсот": amount_words,
            "50000": str(rub),
        },
    )
    wb.save(path)


def order_generate_document(order_id: int) -> tuple[bool, str, Path | None]:
    row = order_get_request(order_id)
    if not row:
        return False, "Заявка не найдена", None

    path = order_document_path(order_id)
    template_path = _order_template_path(str(row.get("doc_type") or ""))
    if not template_path.exists():
        return False, "Не найден шаблон документа ПКО/РКО", None

    shutil.copy2(template_path, path)
    if str(row.get("doc_type") or "").strip().upper() == "ПКО":
        _fill_pko_template(path, row)
    else:
        _fill_rko_template(path, row)
    return True, "Документ сформирован", path


def _yandex_headers() -> dict[str, str]:
    token = os.getenv("YANDEX_DISK_TOKEN", "").strip()
    if token == "":
        return {}
    return {"Authorization": f"OAuth {token}"}


def _safe_disk_name(value: str, fallback: str = "unknown") -> str:
    text = " ".join(str(value or "").strip().split())
    if not text:
        text = fallback
    bad = '\\/:*?"<>|'
    for ch in bad:
        text = text.replace(ch, " ")
    text = "_".join([p for p in text.split(" ") if p])
    return text[:120] if text else fallback


def _yandex_api_json(method: str, endpoint: str, params: dict[str, Any] | None = None) -> tuple[bool, dict[str, Any] | None, str]:
    headers = _yandex_headers()
    if not headers:
        return False, None, "YANDEX_DISK_TOKEN не задан"
    url = "https://cloud-api.yandex.net/v1/disk" + endpoint
    if params:
        url += "?" + urlencode({k: str(v) for k, v in params.items() if v is not None})
    req = urlrequest.Request(url, method=method, headers=headers)
    try:
        with urlrequest.urlopen(req, timeout=20) as resp:
            raw = resp.read().decode("utf-8", "replace").strip()
            data = json.loads(raw) if raw else {}
            return True, data, "OK"
    except urlerror.HTTPError as exc:
        raw = exc.read().decode("utf-8", "replace").strip()
        return False, None, f"HTTP {exc.code}: {raw}"
    except Exception as exc:
        return False, None, str(exc)


def _yandex_api_json_retry(
    method: str, endpoint: str, params: dict[str, Any] | None = None, retries: int = 6
) -> tuple[bool, dict[str, Any] | None, str]:
    last_ok = False
    last_data: dict[str, Any] | None = None
    last_text = "Неизвестная ошибка"
    for attempt in range(retries):
        ok, data, text = _yandex_api_json(method, endpoint, params)
        last_ok, last_data, last_text = ok, data, text
        if ok:
            return ok, data, text
        if "HTTP 423" not in text and "DiskResourceLockedError" not in text and "HTTP 500" not in text:
            return ok, data, text
        time.sleep(0.6 + attempt * 0.4)
    return last_ok, last_data, last_text


def _yandex_mkdirs(remote_dir: str) -> tuple[bool, str]:
    parts = [p for p in str(remote_dir).split("/") if p]
    curr = ""
    for part in parts:
        curr += "/" + part
        ok, _, text = _yandex_api_json_retry("PUT", "/resources", {"path": f"disk:{curr}"})
        if not ok and "HTTP 409" not in text:
            return False, text
    return True, "OK"


def upload_file_to_yandex_disk(local_path: Path, remote_path: str) -> tuple[bool, str, str | None]:
    if not local_path.exists():
        return False, "Файл для загрузки не найден", None
    remote_clean = "/" + "/".join([p for p in str(remote_path).split("/") if p])
    remote_dir = "/".join(remote_clean.split("/")[:-1]) or "/"
    ok_dir, text_dir = _yandex_mkdirs(remote_dir)
    if not ok_dir:
        return False, text_dir, None

    ok_up, data_up, text_up = _yandex_api_json_retry(
        "GET",
        "/resources/upload",
        {"path": f"disk:{remote_clean}", "overwrite": "true"},
    )
    if not ok_up or not data_up or "href" not in data_up:
        return False, f"Не удалось получить ссылку загрузки: {text_up}", None

    try:
        with local_path.open("rb") as f:
            body = f.read()
        upload_req = urlrequest.Request(
            str(data_up["href"]),
            data=body,
            method="PUT",
            headers={"Content-Type": "application/octet-stream"},
        )
        with urlrequest.urlopen(upload_req, timeout=60):
            pass
    except Exception as exc:
        return False, f"Ошибка загрузки файла: {exc}", None

    private_url = f"https://disk.yandex.ru/client/disk{quote(remote_clean)}"
    ok_pub, _, text_pub = _yandex_api_json_retry("PUT", "/resources/publish", {"path": f"disk:{remote_clean}"})
    if not ok_pub and "HTTP 409" not in text_pub:
        if "HTTP 403" in text_pub:
            return True, "Файл загружен в Яндекс.Диск (без публичной ссылки)", private_url
        return False, f"Файл загружен, но не опубликован: {text_pub}", None

    ok_meta, data_meta, text_meta = _yandex_api_json_retry(
        "GET",
        "/resources",
        {"path": f"disk:{remote_clean}", "fields": "public_url"},
    )
    if not ok_meta:
        return True, f"Файл загружен в Яндекс.Диск: {text_meta}", private_url
    public_url = str((data_meta or {}).get("public_url") or "").strip()
    if public_url == "":
        return True, "Файл загружен в Яндекс.Диск (без публичной ссылки)", private_url
    return True, "OK", public_url


def order_upload_document_to_yandex(order_id: int, row: dict[str, Any], local_path: Path) -> tuple[bool, str, str | None]:
    base_dir = os.getenv("YANDEX_DISK_BASE_PATH", "/Infinity").strip() or "/Infinity"
    safe_base = "/" + "/".join([p for p in base_dir.split("/") if p])
    filename = f"{row.get('doc_type', 'ORDER')}_{row.get('doc_number', order_id)}_{order_id}.xlsx"
    remote_path = f"{safe_base}/order-bot/{datetime.now().strftime('%Y-%m')}/{filename}"
    return upload_file_to_yandex_disk(local_path, remote_path)


def docflow_get_application(app_id: int) -> dict[str, Any] | None:
    _ensure_docflow_schema()
    rows = _query(
        db_paths()["docflow"],
        """
        SELECT id, deal_type, contract_no, address, object_type, head_name, agent_name, department_no, status, created_at
        FROM applications
        WHERE id = ?
        LIMIT 1
        """,
        (app_id,),
    )
    return dict(rows[0]) if rows else None


def docflow_upload_bundle_to_yandex(app_id: int) -> tuple[bool, str, str | None]:
    details = docflow_get_application_details(app_id)
    row = docflow_get_application(app_id)
    if not details or not row:
        return False, "Данные заявки не найдены", None
    cached_url = str(details.get("yadisk_url") or "").strip()
    if cached_url and "/client/disk" not in cached_url:
        return True, "OK", cached_url

    doc_path = Path(str(details.get("document_path") or ""))
    uploads_root = docflow_uploads_dir(app_id)
    base_dir = os.getenv("YANDEX_DISK_BASE_PATH", "/Infinity").strip() or "/Infinity"
    safe_base = "/" + "/".join([p for p in base_dir.split("/") if p])
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    folder_name = "_".join(
        [
            _safe_disk_name(str(row.get("agent_name") or ""), "agent"),
            _safe_disk_name(str(row.get("contract_no") or ""), "no-contract"),
            _safe_disk_name(str(row.get("deal_type") or ""), "deal"),
            str(app_id),
            ts,
        ]
    )
    remote_root = f"{safe_base}/docflow/{datetime.now().strftime('%Y-%m')}/{folder_name}"
    private_folder_url = f"https://disk.yandex.ru/client/disk{quote(remote_root)}"

    remote_doc_path = ""
    if doc_path.exists():
        remote_doc_path = f"{remote_root}/{doc_path.name}"
        ok_doc, text_doc, _ = upload_file_to_yandex_disk(doc_path, remote_doc_path)
        if not ok_doc:
            return False, text_doc, None

    if uploads_root.exists():
        for file_path in uploads_root.rglob("*"):
            if not file_path.is_file():
                continue
            rel = file_path.relative_to(uploads_root).as_posix()
            ok_file, text_file, _ = upload_file_to_yandex_disk(file_path, f"{remote_root}/{rel}")
            if not ok_file:
                return False, text_file, None

    ok_pub, _, text_pub = _yandex_api_json_retry("PUT", "/resources/publish", {"path": f"disk:{remote_root}"})
    if not ok_pub and "HTTP 409" not in text_pub and "HTTP 403" not in text_pub:
        return False, f"Папка загружена, но не опубликована: {text_pub}", None
    ok_meta, data_meta, text_meta = _yandex_api_json_retry(
        "GET",
        "/resources",
        {"path": f"disk:{remote_root}", "fields": "public_url"},
    )
    if ok_meta:
        public_url = str((data_meta or {}).get("public_url") or "").strip()
        if public_url:
            _execute(
                db_paths()["docflow"],
                "UPDATE web_application_details SET yadisk_url = ?, updated_at = CURRENT_TIMESTAMP WHERE app_id = ?",
                (public_url, app_id),
            )
            return True, "Пакет документов загружен в Яндекс.Диск", public_url

    # fallback: try public link for protocol document
    if remote_doc_path:
        ok_pub_doc, _, text_pub_doc = _yandex_api_json_retry("PUT", "/resources/publish", {"path": f"disk:{remote_doc_path}"})
        if ok_pub_doc or "HTTP 409" in text_pub_doc:
            ok_meta_doc, data_meta_doc, _ = _yandex_api_json_retry(
                "GET",
                "/resources",
                {"path": f"disk:{remote_doc_path}", "fields": "public_url"},
            )
            if ok_meta_doc:
                public_doc_url = str((data_meta_doc or {}).get("public_url") or "").strip()
                if public_doc_url:
                    _execute(
                        db_paths()["docflow"],
                        "UPDATE web_application_details SET yadisk_url = ?, updated_at = CURRENT_TIMESTAMP WHERE app_id = ?",
                        (public_doc_url, app_id),
                    )
                    return True, "Папка не опубликована, выдана публичная ссылка на протокол", public_doc_url

    # fallback: publish archive of all files
    with tempfile.TemporaryDirectory() as tmp_dir:
        zip_path = Path(tmp_dir) / f"all_documents_{app_id}.zip"
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            if doc_path.exists():
                zf.write(doc_path, arcname=doc_path.name)
            if uploads_root.exists():
                for file_path in uploads_root.rglob("*"):
                    if file_path.is_file():
                        zf.write(file_path, arcname=f"files/{file_path.relative_to(uploads_root).as_posix()}")
        remote_zip_path = f"{remote_root}/all_documents_{app_id}.zip"
        ok_zip, _, _ = upload_file_to_yandex_disk(zip_path, remote_zip_path)
        if ok_zip:
            ok_pub_zip, _, text_pub_zip = _yandex_api_json_retry("PUT", "/resources/publish", {"path": f"disk:{remote_zip_path}"})
            if ok_pub_zip or "HTTP 409" in text_pub_zip:
                ok_meta_zip, data_meta_zip, _ = _yandex_api_json_retry(
                    "GET",
                    "/resources",
                    {"path": f"disk:{remote_zip_path}", "fields": "public_url"},
                )
                if ok_meta_zip:
                    public_zip_url = str((data_meta_zip or {}).get("public_url") or "").strip()
                    if public_zip_url:
                        _execute(
                            db_paths()["docflow"],
                            "UPDATE web_application_details SET yadisk_url = ?, updated_at = CURRENT_TIMESTAMP WHERE app_id = ?",
                            (public_zip_url, app_id),
                        )
                        return True, "Папка не опубликована, выдана публичная ссылка на архив документов", public_zip_url
    if "HTTP 403" in text_pub:
        return (
            False,
            "Папка загружена, но публичная ссылка запрещена токеном Я.Диск (HTTP 403). "
            "Нужно выдать токену право публикации файлов/папок.",
            None,
        )
    return False, f"Папка загружена, но не получена публичная ссылка: {text_meta}", None


def order_pending_requests(department: str = "") -> list[dict[str, Any]]:
    db = db_paths()["order"]
    has_web = bool(_table_columns(db, "web_users"))
    dep = department.strip()
    if has_web:
        sql = """
        SELECT o.id, o.doc_type, o.doc_number, o.user_id, COALESCE(w.email, '') AS email, o.date,
               o.full_name, o.amount, o.status, COALESCE(o.comment, '') AS comment, o.created_at,
               COALESCE(u.department, '') AS department
        FROM orders o
        LEFT JOIN web_users w ON w.telegram_id = CAST(o.user_id AS TEXT)
        LEFT JOIN users u ON u.telegram_id = o.user_id
        WHERE o.status = 'на рассмотрении'
        {dep_clause}
        ORDER BY o.id DESC
        LIMIT 500
        """
        dep_clause = "AND COALESCE(u.department, '') = ?" if dep else ""
        return [dict(r) for r in _query(db, sql.format(dep_clause=dep_clause), (dep,) if dep else ())]
    sql = """
    SELECT o.id, o.doc_type, o.doc_number, o.user_id, '' AS email, o.date, o.full_name, o.amount, o.status,
           COALESCE(o.comment, '') AS comment, o.created_at, COALESCE(u.department, '') AS department
    FROM orders o
    LEFT JOIN users u ON u.telegram_id = o.user_id
    WHERE o.status = 'на рассмотрении'
    {dep_clause}
    ORDER BY o.id DESC
    LIMIT 500
    """
    dep_clause = "AND COALESCE(u.department, '') = ?" if dep else ""
    return [dict(r) for r in _query(db, sql.format(dep_clause=dep_clause), (dep,) if dep else ())]


def order_update_status(order_id: int, status: str, comment: str) -> tuple[bool, str]:
    affected = _execute(
        db_paths()["order"],
        "UPDATE orders SET status = ?, comment = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
        (status, comment, order_id),
    )
    return (affected > 0, "Статус обновлен" if affected > 0 else "Заявка не найдена")


# -------------------------
# Contracts module
# -------------------------

CONTRACT_START_NUMBER = 4988


def _ensure_contracts_schema() -> None:
    db = db_paths()["contracts"]
    _ensure_db_file(db)
    with sqlite3.connect(db) as conn:
        try:
            conn.execute("ALTER TABLE contracts ADD COLUMN signed_date TEXT")
        except sqlite3.OperationalError:
            pass
        conn.commit()


def contracts_register_user(telegram_id: int, full_name: str, department: str) -> tuple[bool, str]:
    _ensure_contracts_schema()
    _execute(
        db_paths()["contracts"],
        "INSERT OR REPLACE INTO users (telegram_id, full_name, department) VALUES (?, ?, ?)",
        (telegram_id, full_name.strip(), department.strip()),
    )
    return True, "Пользователь зарегистрирован"


def contracts_register_web(full_name: str, department: str, email: str, account_password: str) -> tuple[bool, str]:
    telegram_id = int(generate_telegram_id("contracts", email))
    ok, text = contracts_register_user(telegram_id, full_name, department)
    if not ok:
        return ok, text
    ok2, text2 = web_register_user("contracts", str(telegram_id), email, account_password)
    if not ok2:
        return False, text2
    return True, "Пользователь зарегистрирован"


def contracts_get_user(telegram_id: int) -> dict[str, Any] | None:
    _ensure_contracts_schema()
    rows = _query(
        db_paths()["contracts"],
        "SELECT telegram_id, full_name, department FROM users WHERE telegram_id = ? LIMIT 1",
        (telegram_id,),
    )
    return dict(rows[0]) if rows else None


def _next_contract_number(db: Path, department: str) -> str:
    cnt = _query(db, "SELECT COUNT(*) AS total FROM contracts")
    seq = int(cnt[0]["total"]) + CONTRACT_START_NUMBER
    return f"{seq}/{department}"


def _sync_contract_to_google_sheets(payload: dict[str, Any]) -> tuple[bool, str]:
    webhook = os.getenv("GOOGLE_SHEETS_WEBHOOK_URL", "").strip()
    if webhook == "":
        return False, "GOOGLE_SHEETS_WEBHOOK_URL не задан"
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urlrequest.Request(
        webhook,
        data=body,
        headers={"Content-Type": "application/json; charset=utf-8"},
        method="POST",
    )
    try:
        with urlrequest.urlopen(req, timeout=10) as resp:
            status = int(getattr(resp, "status", 200) or 200)
            if 200 <= status < 300:
                return True, "OK"
            return False, f"HTTP {status}"
    except urlerror.HTTPError as exc:
        return False, f"HTTP {exc.code}"
    except Exception as exc:
        return False, str(exc)


def contracts_create(telegram_id: int, form: str, address: str) -> tuple[bool, str]:
    _ensure_contracts_schema()
    db = db_paths()["contracts"]
    user = _query(
        db,
        "SELECT full_name, department FROM users WHERE telegram_id = ? LIMIT 1",
        (telegram_id,),
    )
    if not user:
        return False, "Пользователь не зарегистрирован"
    full_name = str(user[0]["full_name"] or "").strip()
    dep = str(user[0]["department"] or "").strip()
    number = _next_contract_number(db, dep)
    _execute(
        db,
        "INSERT INTO contracts (number, user_id, date, form, address, status) VALUES (?, ?, date('now'), ?, ?, 'Не подписан')",
        (number, telegram_id, form, address),
    )
    created_row = _query(
        db,
        "SELECT id, number, user_id, date, form, address, status FROM contracts WHERE number = ? LIMIT 1",
        (number,),
    )
    sync_ok = False
    sync_text = "Не выполнена"
    if created_row:
        row = dict(created_row[0])
        sync_ok, sync_text = _sync_contract_to_google_sheets(
            {
                "source": "web-admin",
                "event": "contract_created",
                "contract_id": row.get("id"),
                "number": row.get("number"),
                "user_id": row.get("user_id"),
                "full_name": full_name,
                "department": dep,
                "date": row.get("date"),
                "form": row.get("form"),
                "address": row.get("address"),
                "status": row.get("status"),
                "created_at": datetime.utcnow().isoformat(sep=" "),
            }
        )
    if sync_ok:
        return True, f"Договор создан: {number}. Google Sheets: синхронизировано"
    return True, f"Договор создан: {number}. Google Sheets: {sync_text}"


def contracts_list(user_id: int | None = None, only_active: bool = False) -> list[dict[str, Any]]:
    _ensure_contracts_schema()
    db = db_paths()["contracts"]
    sql = "SELECT id, number, user_id, date, form, address, status, COALESCE(signed_date, '') AS signed_date, created_at FROM contracts"
    clauses: list[str] = []
    params: list[Any] = []
    if user_id is not None:
        clauses.append("user_id = ?")
        params.append(user_id)
    if only_active:
        clauses.append("status = 'Не подписан'")
    if clauses:
        sql += " WHERE " + " AND ".join(clauses)
    sql += " ORDER BY id DESC LIMIT 500"
    return [dict(r) for r in _query(db, sql, tuple(params))]


def contracts_update_status(contract_id: int, status: str) -> tuple[bool, str]:
    _ensure_contracts_schema()
    if str(status or "").strip().lower() == "подписан":
        row = _query(db_paths()["contracts"], "SELECT COALESCE(signed_date, '') AS signed_date FROM contracts WHERE id = ? LIMIT 1", (contract_id,))
        if not row:
            return False, "Договор не найден"
        if str(row[0]["signed_date"] or "").strip() == "":
            return False, "Сначала укажите дату подписания"
    affected = _execute(db_paths()["contracts"], "UPDATE contracts SET status = ? WHERE id = ?", (status, contract_id))
    return (affected > 0, "Статус обновлен" if affected > 0 else "Договор не найден")


def contracts_mark_signed_for_user(user_id: int, contract_id: int, signed_date: str) -> tuple[bool, str]:
    _ensure_contracts_schema()
    db = db_paths()["contracts"]
    clean_signed_date = str(signed_date or "").strip()
    if clean_signed_date == "":
        return False, "Укажите дату подписания"
    row = _query(db, "SELECT id, user_id, status FROM contracts WHERE id = ? LIMIT 1", (contract_id,))
    if not row:
        return False, "Договор не найден"
    owner_id = int(row[0]["user_id"] or 0)
    if owner_id != int(user_id):
        return False, "Недостаточно прав для изменения статуса"
    if str(row[0]["status"] or "").strip().lower() == "подписан":
        return True, "Статус уже установлен: Подписан"
    affected = _execute(db, "UPDATE contracts SET status = 'Подписан', signed_date = ? WHERE id = ?", (clean_signed_date, contract_id))
    return (affected > 0, "Статус изменен на Подписан" if affected > 0 else "Не удалось обновить статус")


def contracts_templates() -> list[dict[str, str]]:
    files_dir = BASE / "contract-register" / "files"
    if not files_dir.exists():
        return []
    return [{"name": p.name, "path": str(p)} for p in sorted(files_dir.glob("*.docx"))]


# -------------------------
# Docflow module
# -------------------------

DOCFLOW_PASSWORD = "080323"

DOCFLOW_QUESTION_TEXTS: dict[str, str] = {
    "q1": "Новорожденные дети без регистрации не проживают",
    "q2": "На Объекте незарегистрированная перепланировка",
    "q3": "Данные о незарегистрированной перепланировке в документах БТИ",
    "q4": "Претензии третьих лиц в отношении прав на Объект",
    "q5": "У собственника/пользователя есть признаки неадекватного поведения/псих.заболевания",
    "q6": "Задолженность за электроэнергию/коммунальные платежи/капремонт",
    "q7": "Дом планируется",
    "q8": "Объект перед сделкой занимают",
    "q9": "Объект продается",
    "q10": "К моменту сделки на объекте зарегистрировано ___ человек, из них несовершеннолетних ___",
    "q11": "Срок владения Объектом (ручной ввод)",
    "q12": "Является единственным жильем на момент продажи",
    "q13": "Заявление о личном участии в сделке",
    "q14": "Средства материнского капитала на приобретение Объекта",
    "q15": "Относится ли Объект к объектам культурного наследия",
}

DOCFLOW_QUESTION_OPTIONS: dict[str, list[str]] = {
    "q1": ["не проживают", "проживают"],
    "q2": ["отсутствует", "имеется"],
    "q3": ["нет", "есть"],
    "q4": ["отсутствуют", "имеются"],
    "q5": ["нет", "да"],
    "q6": ["отсутствует", "имеется"],
    "q7": ["не планируется", "под снос"],
    "q8": ["физически свободен", "собственники", "наниматели"],
    "q9": ["лично собственником", "по доверенности"],
    "q12": ["нет", "да"],
    "q13": ["не было", "было"],
    "q14": ["не использовались", "использовались"],
    "q15": ["нет", "да"],
}

DOCFLOW_UPLOAD_CATEGORIES: dict[str, str] = {
    "passport": "Паспорт",
    "egrn": "ЕГРН",
    "lawyer_task": "Задания от юриста",
    "other": "Прочее",
}


def docflow_questionnaire() -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for i in range(1, 16):
        key = f"q{i}"
        items.append(
            {
                "key": key,
                "title": DOCFLOW_QUESTION_TEXTS.get(key, f"Вопрос {i}"),
                "options": DOCFLOW_QUESTION_OPTIONS.get(key, []),
                "free_text": key not in DOCFLOW_QUESTION_OPTIONS,
            }
        )
    return items


def docflow_upload_category_map() -> dict[str, str]:
    return dict(DOCFLOW_UPLOAD_CATEGORIES)


def docflow_register_user(telegram_id: str, password: str, full_name: str, department_no: str) -> tuple[bool, str]:
    _ensure_docflow_schema()
    if password != DOCFLOW_PASSWORD:
        return False, "Неверный пароль регистрации"
    db = db_paths()["docflow"]
    cols = _table_columns(db, "users")
    values: dict[str, Any] = {
        "telegram_id": telegram_id,
        "full_name": full_name.strip(),
        "department_no": department_no.strip(),
        "role": "agent",
    }
    if "is_active" in cols:
        values["is_active"] = 0
    if "is_approved" in cols:
        values["is_approved"] = 0
    if "created_at" in cols:
        values["created_at"] = datetime.utcnow().isoformat(sep=" ")
    fields = [f for f in values.keys() if f in cols]
    sql = f"INSERT OR REPLACE INTO users ({', '.join(fields)}) VALUES ({', '.join(['?'] * len(fields))})"
    _execute(db, sql, tuple(values[f] for f in fields))
    return True, "Пользователь зарегистрирован и отправлен на подтверждение РОП"


def docflow_register_web(
    access_password: str, full_name: str, department_no: str, email: str, account_password: str
) -> tuple[bool, str]:
    _ensure_docflow_schema()
    telegram_id = generate_telegram_id("docflow", email)
    ok, text = docflow_register_user(telegram_id, access_password, full_name, department_no)
    if not ok:
        return ok, text
    ok2, text2 = web_register_user("docflow", telegram_id, email, account_password)
    if not ok2:
        return False, text2
    return True, "Пользователь зарегистрирован и отправлен на подтверждение РОП"


def docflow_get_user(telegram_id: str) -> dict[str, Any] | None:
    _ensure_docflow_schema()
    db = db_paths()["docflow"]
    cols = _table_columns(db, "users")
    sql = "SELECT telegram_id, full_name, department_no, role"
    if "is_active" in cols:
        sql += ", is_active"
    if "is_approved" in cols:
        sql += ", is_approved"
    sql += " FROM users WHERE telegram_id = ? LIMIT 1"
    rows = _query(db, sql, (telegram_id,))
    return dict(rows[0]) if rows else None


def _legacy_docflow_pending_users() -> list[dict[str, Any]]:
    _ensure_docflow_schema()
    db = db_paths()["docflow"]
    cols = _table_columns(db, "users")
    has_web = bool(_table_columns(db, "web_users"))
    select_email = "COALESCE(w.email, '') AS email," if has_web else "'' AS email,"
    join_web = "LEFT JOIN web_users w ON w.telegram_id = u.telegram_id" if has_web else ""
    sql = f"SELECT {select_email} u.telegram_id, u.full_name, u.department_no, u.role"
    if "is_approved" in cols:
        sql += ", u.is_approved"
    if "is_active" in cols:
        sql += ", u.is_active"
    sql += " FROM users u "
    if join_web:
        sql += join_web + " "
    if "is_approved" in cols:
        sql += "WHERE COALESCE(u.is_approved, 0) = 0 "
    sql += "ORDER BY u.telegram_id DESC LIMIT 300"
    return [dict(r) for r in _query(db, sql)]


def _legacy_docflow_approve_user(telegram_id: str, approve: bool) -> tuple[bool, str]:
    _ensure_docflow_schema()
    db = db_paths()["docflow"]
    cols = _table_columns(db, "users")
    if "is_approved" in cols and "is_active" in cols:
        affected = _execute(
            db,
            "UPDATE users SET is_approved = ?, is_active = ? WHERE telegram_id = ?",
            (1 if approve else 0, 1 if approve else 0, telegram_id),
        )
    elif "is_active" in cols:
        affected = _execute(db, "UPDATE users SET is_active = ? WHERE telegram_id = ?", (1 if approve else 0, telegram_id))
    else:
        affected = _execute(db, "UPDATE users SET role = role WHERE telegram_id = ?", (telegram_id,))
    return (affected > 0, "Статус подтверждения обновлен" if affected > 0 else "Пользователь не найден")


def docflow_create_application(
    agent_telegram_id: str, deal_type: str, contract_no: str, address: str, object_type: str, head_name: str
) -> tuple[bool, str]:
    ok, text, _ = docflow_create_application_full(
        agent_telegram_id=agent_telegram_id,
        deal_type=deal_type,
        contract_no=contract_no,
        address=address,
        object_type=object_type,
        head_name=head_name,
    )
    return ok, text


def docflow_create_application_full(
    agent_telegram_id: str,
    deal_type: str,
    contract_no: str,
    address: str,
    object_type: str,
    head_name: str,
) -> tuple[bool, str, int | None]:
    _ensure_docflow_schema()
    db = db_paths()["docflow"]
    user_cols = _table_columns(db, "users")
    user_select = "full_name"
    if "id" in user_cols:
        user_select = "id, full_name"
    user = _query(db, f"SELECT {user_select} FROM users WHERE telegram_id = ? LIMIT 1", (agent_telegram_id,))
    if not user:
        return False, "Сотрудник не найден", None
    agent_id = user[0]["id"] if "id" in user[0].keys() else None
    agent_name = str(user[0]["full_name"] or "").strip()
    user_dep_rows = _query(db, "SELECT department_no FROM users WHERE telegram_id = ? LIMIT 1", (agent_telegram_id,))
    department_no = str(user_dep_rows[0]["department_no"] or "").strip() if user_dep_rows else ""
    if not agent_name:
        return False, "Не заполнено ФИО сотрудника в базе", None

    cols = _table_columns(db, "applications")
    values: dict[str, Any] = {
        "deal_type": deal_type,
        "contract_no": contract_no,
        "address": address,
        "object_type": object_type,
        "head_name": head_name,
        "agent_name": agent_name,
        "department_no": department_no,
        "status": "CREATED",
        "created_at": datetime.utcnow().isoformat(sep=" "),
        "updated_at": datetime.utcnow().isoformat(sep=" "),
    }
    if agent_id is not None and "agent_id" in cols:
        values["agent_id"] = agent_id
    fields = [f for f in values.keys() if f in cols]
    if not fields:
        return False, "Таблица applications не содержит ожидаемых полей", None
    sql = f"INSERT INTO applications ({', '.join(fields)}) VALUES ({', '.join(['?'] * len(fields))})"
    try:
        with sqlite3.connect(db) as conn:
            cur = conn.execute(sql, tuple(values[f] for f in fields))
            conn.commit()
            app_id = int(cur.lastrowid or 0)
            if app_id <= 0:
                last = conn.execute("SELECT id FROM applications ORDER BY id DESC LIMIT 1").fetchone()
                app_id = int(last[0]) if last else 0
        return True, "Заявка создана", app_id if app_id > 0 else None
    except sqlite3.Error as exc:
        return False, f"Ошибка создания заявки: {exc}", None


def _docflow_details_db_table() -> None:
    db = db_paths()["docflow"]
    _ensure_db_file(db)
    with sqlite3.connect(db) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS web_application_details (
                app_id INTEGER PRIMARY KEY,
                answers_json TEXT NOT NULL,
                document_path TEXT NOT NULL,
                uploads_json TEXT NOT NULL,
                yadisk_url TEXT NOT NULL DEFAULT '',
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        try:
            conn.execute("ALTER TABLE web_application_details ADD COLUMN yadisk_url TEXT NOT NULL DEFAULT ''")
        except sqlite3.OperationalError:
            pass
        conn.commit()


def docflow_document_path(app_id: int) -> Path:
    docs_dir = BASE / "doc-flow-bot" / "app" / "web_documents"
    docs_dir.mkdir(parents=True, exist_ok=True)
    return docs_dir / f"application_{app_id}.docx"


def docflow_uploads_dir(app_id: int) -> Path:
    uploads_dir = BASE / "doc-flow-bot" / "app" / "web_uploads" / f"application_{app_id}"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    return uploads_dir


def _a(answers: dict[str, str], key: str, default: str = "") -> str:
    return str(answers.get(key, "")).strip() or default


def docflow_generate_application_document(
    app_id: int,
    app_row: dict[str, Any],
    answers: dict[str, str],
    uploaded_files: Any,
) -> Path:
    from docx import Document

    path = docflow_document_path(app_id)
    doc = Document()
    doc.add_heading("ПРОТОКОЛ ПРОВЕРКИ ОБЪЕКТА", 0)
    doc.add_paragraph(f"Заявка № {app_id}")
    doc.add_paragraph(f"Тип сделки: {app_row.get('deal_type', '')}")
    doc.add_paragraph(f"№ договора: {app_row.get('contract_no', '')}")
    doc.add_paragraph(f"Адрес: {app_row.get('address', '')}")
    doc.add_paragraph(f"Сотрудник: {app_row.get('agent_name', '')}")
    doc.add_paragraph("")
    doc.add_paragraph("5. Обстоятельства, подлежащие проверке (обязательно для заполнения):")
    doc.add_paragraph("Заполняется Сотрудником:")
    doc.add_paragraph(f"1. Новорожденные дети без регистрации {_a(answers, 'q1', 'не проживают')}")
    doc.add_paragraph(f"2. На Объекте незарегистрированная перепланировка {_a(answers, 'q2', 'отсутствует')}")
    doc.add_paragraph(f"3. Данные о незарегистрированной перепланировке {_a(answers, 'q3', 'нет')} в документах БТИ")
    doc.add_paragraph(f"4. Претензии третьих лиц в отношении прав на Объект {_a(answers, 'q4', 'отсутствуют')}")
    doc.add_paragraph(
        f"5. У собств./польз. есть признаки неадекватного поведения/ псих.заболевания {_a(answers, 'q5', 'нет')}"
    )
    doc.add_paragraph(f"6. Задолженность за электроэнергию/коммунальные платежи/капремонт {_a(answers, 'q6', 'нет данных')}")
    doc.add_paragraph(f"7. Дом, планируется {_a(answers, 'q7', 'не планируется')}")
    doc.add_paragraph(f"8. Объект перед сделкой занимают {_a(answers, 'q8', 'физически свободен')}")
    doc.add_paragraph(f"9. Объект продается {_a(answers, 'q9', 'лично собственником')}")
    doc.add_paragraph(f"10. К моменту сделки на объекте зарегистрировано {_a(answers, 'q10', '0')}")
    doc.add_paragraph(f"11. Срок владения Объектом - {_a(answers, 'q11', '-')}")
    doc.add_paragraph(f"12. Является единственным жильем на момент продажи: {_a(answers, 'q12', 'нет')}")
    doc.add_paragraph(f"13. Заявление о личном участии в сделке: {_a(answers, 'q13', 'не было')}")
    doc.add_paragraph(f"14. Средства материнского капитала на приобретение Объекта: {_a(answers, 'q14', 'не использовались')}")
    doc.add_paragraph(f"15. Относится ли Объект к объектам культурного наследия: {_a(answers, 'q15', 'нет')}")
    doc.add_paragraph("Факты установил: _________________________________")
    doc.add_paragraph("подпись сотрудника")
    doc.add_paragraph("")
    doc.add_paragraph("Заполняется юристом:")
    doc.add_paragraph("Сделки, совершенные ранее по доверенности проводились/не проводились/не установлено")
    doc.add_paragraph("Обременения/ограничения права собственности имеются/отсутствуют")
    doc.add_paragraph("Аресты запрещения на объект накладывались/ не накладывались / не установлено")
    doc.add_paragraph("Судебные споры в отношении Объекта имеются/отсутствуют")
    doc.add_paragraph("Сделки, совершенные на коротком отрезке времени проводились / не проводились / не установлено")
    doc.add_paragraph("Исполнительные производства в отношении собственника имеются/отсутствуют/не установлено")
    doc.add_paragraph("Сведения о недействительности паспортов участников имеются/отсутствуют")
    doc.add_paragraph("Сведения о банкротстве/признаках неплатежеспособности собственника имеются/отсутствуют")
    doc.add_paragraph(
        "Лица, временно снятые с регистрационного учета присутствуют / отсутствуют / не установлено "
        "(военная служба, заключение, безвестное отсутствие, дом престарелых, интернат)"
    )
    doc.add_paragraph("Собственники/пользователи/супруги в ПНД / НД состоят / не состоят / не установлено")
    doc.add_paragraph("Пункты 12,13,14,15: проверены")
    doc.add_paragraph("Дополнительные сведения:")
    doc.add_paragraph("Подпись юриста:_______________")
    doc.add_paragraph("РАЗРЕШЕНИЕ СДЕЛКИ ПО ОБЪЕКТУ:")
    doc.add_paragraph("Юридический отдел: Разрешено / Не разрешено / Разрешено с условиями:")
    doc.add_paragraph("___________________________________________________________________________________________")
    doc.add_paragraph("")
    doc.add_heading("Загруженные документы", level=1)
    if isinstance(uploaded_files, dict):
        has_files = False
        for category_key, folder_name in DOCFLOW_UPLOAD_CATEGORIES.items():
            files = uploaded_files.get(category_key, [])
            if not files:
                continue
            has_files = True
            doc.add_paragraph(f"{folder_name}:")
            for file_name in files:
                doc.add_paragraph(f" - {file_name}")
        if not has_files:
            doc.add_paragraph("Файлы не загружены")
    elif isinstance(uploaded_files, list) and uploaded_files:
        for file_name in uploaded_files:
            doc.add_paragraph(str(file_name))
    else:
        doc.add_paragraph("Файлы не загружены")
    doc.save(path)
    return path


def docflow_save_application_details(
    app_id: int, answers: dict[str, str], document_path: Path, uploads: Any, yadisk_url: str = ""
) -> tuple[bool, str]:
    _docflow_details_db_table()
    db = db_paths()["docflow"]
    try:
        with sqlite3.connect(db) as conn:
            conn.execute(
                """
                INSERT INTO web_application_details (app_id, answers_json, document_path, uploads_json, updated_at)
                VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
                ON CONFLICT(app_id) DO UPDATE SET
                    answers_json = excluded.answers_json,
                    document_path = excluded.document_path,
                    uploads_json = excluded.uploads_json,
                    updated_at = CURRENT_TIMESTAMP
                """,
                (app_id, json.dumps(answers, ensure_ascii=False), str(document_path), json.dumps(uploads, ensure_ascii=False)),
            )
            if str(yadisk_url or "").strip():
                conn.execute(
                    "UPDATE web_application_details SET yadisk_url = ?, updated_at = CURRENT_TIMESTAMP WHERE app_id = ?",
                    (str(yadisk_url).strip(), app_id),
                )
            conn.commit()
        return True, "OK"
    except sqlite3.Error as exc:
        return False, str(exc)


def docflow_get_application_details(app_id: int) -> dict[str, Any] | None:
    _docflow_details_db_table()
    db = db_paths()["docflow"]
    rows = _query(
        db,
        "SELECT app_id, answers_json, document_path, uploads_json, COALESCE(yadisk_url, '') AS yadisk_url, created_at, updated_at FROM web_application_details WHERE app_id = ? LIMIT 1",
        (app_id,),
    )
    if not rows:
        return None
    row = dict(rows[0])
    return {
        "app_id": row.get("app_id"),
        "answers": json.loads(str(row.get("answers_json") or "{}")),
        "document_path": str(row.get("document_path") or ""),
        "uploads": json.loads(str(row.get("uploads_json") or "[]")),
        "yadisk_url": str(row.get("yadisk_url") or ""),
        "created_at": row.get("created_at"),
        "updated_at": row.get("updated_at"),
    }


def _ensure_docflow_exchange_schema() -> None:
    _ensure_docflow_schema()
    db = db_paths()["docflow"]
    with sqlite3.connect(db) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS web_application_events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                app_id INTEGER NOT NULL,
                actor_telegram_id TEXT NOT NULL DEFAULT '',
                actor_name TEXT NOT NULL DEFAULT '',
                actor_role TEXT NOT NULL DEFAULT '',
                event_type TEXT NOT NULL DEFAULT 'COMMENT',
                message TEXT NOT NULL DEFAULT '',
                file_paths_json TEXT NOT NULL DEFAULT '[]',
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS web_user_notifications (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_telegram_id TEXT NOT NULL,
                category TEXT NOT NULL DEFAULT 'docflow',
                title TEXT NOT NULL,
                message TEXT NOT NULL,
                link TEXT NOT NULL DEFAULT '',
                is_read INTEGER NOT NULL DEFAULT 0,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        conn.execute("CREATE INDEX IF NOT EXISTS idx_web_app_events_app_id ON web_application_events(app_id)")
        conn.execute(
            "CREATE INDEX IF NOT EXISTS idx_web_user_notifications_user_id ON web_user_notifications(user_telegram_id, is_read)"
        )
        conn.commit()


def docflow_get_application(app_id: int) -> dict[str, Any] | None:
    _ensure_docflow_schema()
    rows = _query(
        db_paths()["docflow"],
        """
        SELECT id, agent_id, agent_name, department_no, deal_type, contract_no, status, created_at
        FROM applications
        WHERE id = ?
        LIMIT 1
        """,
        (app_id,),
    )
    return dict(rows[0]) if rows else None


def docflow_get_agent_telegram_id(app_id: int) -> str:
    app_row = docflow_get_application(app_id)
    if not app_row:
        return ""
    db = db_paths()["docflow"]
    user_cols = _table_columns(db, "users")
    try:
        if "id" in user_cols and app_row.get("agent_id") not in {None, ""}:
            rows = _query(db, "SELECT telegram_id FROM users WHERE id = ? LIMIT 1", (int(app_row["agent_id"]),))
            if rows:
                return str(rows[0]["telegram_id"] or "")
    except (TypeError, ValueError):
        pass
    agent_name = str(app_row.get("agent_name") or "").strip()
    if agent_name:
        rows = _query(db, "SELECT telegram_id FROM users WHERE full_name = ? ORDER BY id DESC LIMIT 1", (agent_name,))
        if rows:
            return str(rows[0]["telegram_id"] or "")
    return ""


def docflow_add_event(
    app_id: int,
    actor_telegram_id: str,
    actor_name: str,
    actor_role: str,
    event_type: str,
    message: str,
    file_paths: list[str] | None = None,
) -> tuple[bool, str]:
    _ensure_docflow_exchange_schema()
    app_row = docflow_get_application(app_id)
    if not app_row:
        return False, "Заявка не найдена"
    event_type_clean = str(event_type or "COMMENT").strip().upper()
    message_clean = str(message or "").strip()
    files_clean = [str(p).strip() for p in (file_paths or []) if str(p).strip()]
    if message_clean == "" and not files_clean:
        return False, "Сообщение или файл обязательны"
    affected = _execute(
        db_paths()["docflow"],
        """
        INSERT INTO web_application_events (app_id, actor_telegram_id, actor_name, actor_role, event_type, message, file_paths_json)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (
            app_id,
            str(actor_telegram_id or "").strip(),
            str(actor_name or "").strip(),
            str(actor_role or "").strip().lower(),
            event_type_clean,
            message_clean,
            json.dumps(files_clean, ensure_ascii=False),
        ),
    )
    return (affected > 0, "Событие сохранено" if affected > 0 else "Не удалось сохранить событие")


def docflow_events(app_id: int, limit: int = 300) -> list[dict[str, Any]]:
    _ensure_docflow_exchange_schema()
    rows = _query(
        db_paths()["docflow"],
        """
        SELECT id, app_id, actor_telegram_id, actor_name, actor_role, event_type, message, file_paths_json, created_at
        FROM web_application_events
        WHERE app_id = ?
        ORDER BY id ASC
        LIMIT ?
        """,
        (app_id, max(int(limit), 1)),
    )
    result: list[dict[str, Any]] = []
    for row in rows:
        item = dict(row)
        try:
            item["file_paths"] = json.loads(str(item.get("file_paths_json") or "[]"))
        except json.JSONDecodeError:
            item["file_paths"] = []
        result.append(item)
    return result


def docflow_add_user_notification(
    user_telegram_id: str,
    title: str,
    message: str,
    link: str = "",
    category: str = "docflow",
) -> None:
    _ensure_docflow_exchange_schema()
    _execute(
        db_paths()["docflow"],
        """
        INSERT INTO web_user_notifications (user_telegram_id, category, title, message, link, is_read)
        VALUES (?, ?, ?, ?, ?, 0)
        """,
        (
            str(user_telegram_id or "").strip(),
            str(category or "docflow").strip(),
            str(title or "").strip(),
            str(message or "").strip(),
            str(link or "").strip(),
        ),
    )


def docflow_user_notifications(user_telegram_id: str, limit: int = 50) -> list[dict[str, Any]]:
    _ensure_docflow_exchange_schema()
    rows = _query(
        db_paths()["docflow"],
        """
        SELECT id, category, title, message, link, is_read, created_at
        FROM web_user_notifications
        WHERE user_telegram_id = ?
        ORDER BY id DESC
        LIMIT ?
        """,
        (str(user_telegram_id or "").strip(), max(int(limit), 1)),
    )
    return [dict(r) for r in rows]


def docflow_mark_user_notification_read(user_telegram_id: str, notification_id: int) -> bool:
    _ensure_docflow_exchange_schema()
    affected = _execute(
        db_paths()["docflow"],
        "UPDATE web_user_notifications SET is_read = 1 WHERE id = ? AND user_telegram_id = ?",
        (notification_id, str(user_telegram_id or "").strip()),
    )
    return affected > 0


def docflow_applications_with_document_link(
    all_rows: bool = False, agent_telegram_id: str = "", department_no: str = ""
) -> list[dict[str, Any]]:
    apps = docflow_applications(department_no=department_no) if all_rows else docflow_applications_by_user(
        agent_telegram_id, department_no=department_no
    )
    result: list[dict[str, Any]] = []
    for app in apps:
        row = dict(app)
        details = docflow_get_application_details(int(row.get("id") or 0))
        row["has_document"] = bool(details and str(details.get("document_path") or "").strip())
        result.append(row)
    return result


def docflow_applications(department_no: str = "") -> list[dict[str, Any]]:
    _ensure_docflow_schema()
    dep = department_no.strip()
    sql = """
    SELECT id, deal_type, contract_no, agent_name, department_no, status, created_at
    FROM applications
    {where_clause}
    ORDER BY id DESC
    LIMIT 500
    """
    where_clause = "WHERE COALESCE(department_no, '') = ?" if dep else ""
    rows = _query(db_paths()["docflow"], sql.format(where_clause=where_clause), (dep,) if dep else ())
    return [dict(r) for r in rows]


def docflow_applications_by_user(agent_telegram_id: str, department_no: str = "") -> list[dict[str, Any]]:
    _ensure_docflow_schema()
    db = db_paths()["docflow"]
    dep = department_no.strip()
    user_cols = _table_columns(db, "users")
    select_cols = "full_name"
    if "id" in user_cols:
        select_cols = "id, full_name"
    user = _query(db, f"SELECT {select_cols} FROM users WHERE telegram_id = ? LIMIT 1", (agent_telegram_id,))
    if not user:
        return []
    user_id = user[0]["id"] if "id" in user[0].keys() else None
    full_name = str(user[0]["full_name"])
    if user_id is not None:
        rows = _query(
            db,
            """
            SELECT id, deal_type, contract_no, agent_name, department_no, status, created_at
            FROM applications
            WHERE agent_id = ? OR agent_name = ?
            ORDER BY id DESC
            LIMIT 500
            """,
            (user_id, full_name),
        )
    else:
        rows = _query(
            db,
            """
            SELECT id, deal_type, contract_no, agent_name, department_no, status, created_at
            FROM applications
            WHERE agent_name = ?
            ORDER BY id DESC
            LIMIT 500
            """,
            (full_name,),
        )
    if dep:
        rows = [r for r in rows if str(r["department_no"] or "").strip() == dep]
    return [dict(r) for r in rows]


def docflow_update_status(app_id: int, status: str, department_no: str = "") -> tuple[bool, str]:
    _ensure_docflow_schema()
    dep = department_no.strip()
    if dep:
        affected = _execute(
            db_paths()["docflow"],
            "UPDATE applications SET status = ? WHERE id = ? AND COALESCE(department_no, '') = ?",
            (status, app_id, dep),
        )
    else:
        affected = _execute(db_paths()["docflow"], "UPDATE applications SET status = ? WHERE id = ?", (status, app_id))
    return (affected > 0, "Статус заявки обновлен" if affected > 0 else "Заявка не найдена")


def docflow_pending_users(department_no: str = "") -> list[dict[str, Any]]:
    _ensure_docflow_schema()
    db = db_paths()["docflow"]
    cols = _table_columns(db, "users")
    has_web = bool(_table_columns(db, "web_users"))
    dep = department_no.strip()
    select_email = "COALESCE(w.email, '') AS email," if has_web else "'' AS email,"
    join_web = "LEFT JOIN web_users w ON w.telegram_id = u.telegram_id" if has_web else ""
    sql = f"SELECT {select_email} u.telegram_id, u.full_name, u.department_no, u.role"
    if "is_approved" in cols:
        sql += ", u.is_approved"
    if "is_active" in cols:
        sql += ", u.is_active"
    sql += " FROM users u "
    if join_web:
        sql += join_web + " "
    where: list[str] = []
    params: list[Any] = []
    if "is_approved" in cols:
        where.append("COALESCE(u.is_approved, 0) = 0")
    if dep:
        where.append("COALESCE(u.department_no, '') = ?")
        params.append(dep)
    if where:
        sql += "WHERE " + " AND ".join(where) + " "
    sql += "ORDER BY u.telegram_id DESC LIMIT 300"
    return [dict(r) for r in _query(db, sql, tuple(params))]


def docflow_approve_user(telegram_id: str, approve: bool, department_no: str = "") -> tuple[bool, str]:
    _ensure_docflow_schema()
    db = db_paths()["docflow"]
    cols = _table_columns(db, "users")
    dep = department_no.strip()
    dep_clause = " AND COALESCE(department_no, '') = ?" if dep else ""
    dep_params: tuple[Any, ...] = (dep,) if dep else ()
    if "is_approved" in cols and "is_active" in cols:
        affected = _execute(
            db,
            f"UPDATE users SET is_approved = ?, is_active = ? WHERE telegram_id = ?{dep_clause}",
            (1 if approve else 0, 1 if approve else 0, telegram_id, *dep_params),
        )
    elif "is_active" in cols:
        affected = _execute(
            db, f"UPDATE users SET is_active = ? WHERE telegram_id = ?{dep_clause}", (1 if approve else 0, telegram_id, *dep_params)
        )
    else:
        affected = _execute(db, f"UPDATE users SET role = role WHERE telegram_id = ?{dep_clause}", (telegram_id, *dep_params))
    return (affected > 0, "Статус подтверждения обновлен" if affected > 0 else "Пользователь не найден")

